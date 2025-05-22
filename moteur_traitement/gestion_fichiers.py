# moteur_traitement/gestion_fichiers.py
import os
import re
from docx import Document # Pour lire et écrire les .docx
import fitz  # PyMuPDF, pour lire les .pdf
from fpdf import FPDF # Pour écrire des .pdf simples

# --- Fonctions de Lecture ---

def lire_contenu_txt(chemin_fichier_txt):
    """Lit le contenu d'un fichier .txt."""
    try:
        with open(chemin_fichier_txt, 'r', encoding='utf-8') as f_in:
            return f_in.read()
    except Exception as e:
        raise ValueError(f"Impossible de lire le fichier TXT '{chemin_fichier_txt}': {e}")

def lire_contenu_docx_texte_complet(chemin_fichier_docx):
    """Lit tout le texte d'un DOCX (paragraphes et tableaux) en une seule chaîne."""
    try:
        doc = Document(chemin_fichier_docx)
        contenu_complet = []
        for para in doc.paragraphs:
            contenu_complet.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para_in_cell in cell.paragraphs:
                        contenu_complet.append(para_in_cell.text)
        return "\n".join(contenu_complet)
    except Exception as e:
        raise ValueError(f"Impossible de lire le contenu texte complet du fichier DOCX '{chemin_fichier_docx}': {e}")

def lire_contenu_pdf_texte_complet(chemin_fichier_pdf):
    """Lit tout le texte d'un PDF page par page."""
    try:
        texte_complet = ""
        with fitz.open(chemin_fichier_pdf) as doc_pdf:
            for page_num in range(len(doc_pdf)):
                page = doc_pdf.load_page(page_num)
                texte_complet += page.get_text("text")
                if page_num < len(doc_pdf) - 1:
                    texte_complet += "\n" 
        return texte_complet.strip()
    except Exception as e:
        raise ValueError(f"Impossible de lire le fichier PDF '{chemin_fichier_pdf}': {e}")

def normaliser_texte(texte_brut):
    """Normalise les sauts de ligne, les espaces multiples, et supprime les espaces en début/fin de chaque ligne."""
    lignes_normalisees = []
    if texte_brut is None: 
        return ""
    for ligne in texte_brut.splitlines():
        ligne_sans_espaces_multiples = re.sub(r'[ \t]+', ' ', ligne)
        ligne_traitee = ligne_sans_espaces_multiples.strip()
        lignes_normalisees.append(ligne_traitee)
    return "\n".join(lignes_normalisees)

# --- Fonctions d'Écriture ---

def sauvegarder_fichier_txt(chemin_fichier_sortie, contenu_texte):
    """Sauvegarde du contenu texte dans un fichier .txt."""
    try:
        dossier_sortie = os.path.dirname(chemin_fichier_sortie)
        if dossier_sortie and not os.path.exists(dossier_sortie):
            os.makedirs(dossier_sortie)
        with open(chemin_fichier_sortie, 'w', encoding='utf-8') as f_out:
            f_out.write(contenu_texte)
        print(f"Fichier TXT sauvegardé sous : {chemin_fichier_sortie}")
    except Exception as e:
        raise IOError(f"Erreur lors de l'écriture du fichier TXT '{chemin_fichier_sortie}': {e}")

def sauvegarder_document_docx_simple(chemin_fichier_sortie, contenu_texte_pseudonymise):
    """Crée un nouveau DOCX et y écrit le texte paragraphe par paragraphe (formatage simple)."""
    try:
        dossier_sortie = os.path.dirname(chemin_fichier_sortie)
        if dossier_sortie and not os.path.exists(dossier_sortie):
            os.makedirs(dossier_sortie)

        doc_sortie = Document()
        if contenu_texte_pseudonymise: # S'assurer qu'il y a du texte à écrire
            for paragraphe_texte in contenu_texte_pseudonymise.splitlines():
                doc_sortie.add_paragraph(paragraphe_texte)
        else: # Ajouter un paragraphe vide si le contenu est vide pour créer un docx valide
            doc_sortie.add_paragraph("")
        doc_sortie.save(chemin_fichier_sortie)
        print(f"Fichier DOCX (simple) sauvegardé sous : {chemin_fichier_sortie}")
    except Exception as e:
        raise IOError(f"Erreur lors de l'écriture du fichier DOCX (simple) '{chemin_fichier_sortie}': {e}")

def sauvegarder_document_pdf_simple(chemin_fichier_sortie, contenu_texte):
    """Crée un nouveau PDF et y écrit le texte (formatage simple)."""
    try:
        dossier_sortie = os.path.dirname(chemin_fichier_sortie)
        if dossier_sortie and not os.path.exists(dossier_sortie):
            os.makedirs(dossier_sortie)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)

        # Tentative d'utiliser une police Unicode (DejaVu)
        # Assurez-vous que le fichier .ttf est accessible (ex: dans le même dossier ou un dossier 'fonts')
        # ou installez la police sur votre système et FPDF2 la trouvera peut-être.
        # Pour un POC, si DejaVu n'est pas trouvée, on se rabat sur Arial.
        try:
            # Vous devrez peut-être télécharger DejaVuSansCondensed.ttf et le placer dans un dossier
            # 'fonts' à côté de ce script, puis utiliser un chemin comme :
            # font_path = os.path.join(os.path.dirname(__file__), 'fonts', 'DejaVuSansCondensed.ttf')
            # pdf.add_font('DejaVu', '', font_path, uni=True)
            # Pour l'instant, on laisse fpdf2 chercher. S'il échoue, il passera à Arial.
            #pdf.add_font('DejaVu', '', 'DejaVuSansCondensed.ttf', uni=True) 
            font_path = os.path.join(os.path.dirname(__file__), 'fonts', 'DejaVuSansCondensed.ttf')
            pdf.add_font('DejaVu', '', font_path, uni=True)
            pdf.set_font("DejaVu", size=10)
            print("Utilisation de la police DejaVu pour le PDF.")
        except RuntimeError:
            print("Avertissement: Police DejaVu (DejaVuSansCondensed.ttf) non trouvée ou erreur de chargement. Utilisation d'Arial (peut mal gérer certains caractères UTF-8).")
            pdf.set_font("Arial", size=10)

        # Écrire le texte ligne par ligne
        lignes_a_ecrire = contenu_texte.splitlines()
        if not lignes_a_ecrire and contenu_texte: # Si le texte n'est pas vide mais n'a pas de \n (une seule ligne)
            lignes_a_ecrire = [contenu_texte]
        elif not lignes_a_ecrire and not contenu_texte: # Si le texte est complètement vide
             lignes_a_ecrire = [""] # Pour créer un PDF avec au moins une page vide


        for line in lignes_a_ecrire:
            try:
                # La méthode write() de FPDF2 gère mieux les caractères spéciaux si la police est Unicode.
                # Pour les polices standard comme Arial, l'encodage est plus délicat.
                pdf.write(5, line) # 5 est la hauteur de ligne suggérée
                pdf.ln() # Passage à la ligne suivante dans le PDF
            except UnicodeEncodeError:
                # Solution de repli si la police actuelle (probablement Arial) ne gère pas le caractère
                cleaned_line = line.encode('latin-1', 'replace').decode('latin-1')
                pdf.write(5, cleaned_line)
                pdf.ln()
            except Exception as e_write_line:
                print(f"Erreur lors de l'écriture d'une ligne dans le PDF ('{line[:30]}...'): {e_write_line}")
                pdf.write(5, "[ligne avec erreur d'encodage]") # Indiquer une erreur dans le PDF
                pdf.ln()

        pdf.output(chemin_fichier_sortie, "F")
        print(f"Fichier PDF (simple, texte seulement) sauvegardé sous : {chemin_fichier_sortie}")
    except Exception as e:
        raise IOError(f"Erreur lors de l'écriture du fichier PDF (simple) '{chemin_fichier_sortie}': {e}")